from pdf2docx import Converter
from PIL import Image
import fitz
import openpyxl
import ebooklib.epub as epub
from typing import List
from docx import Document
from docx.shared import Inches
import io

def convert_to_word(pdf_file: str) -> str:
    output_file = "output.docx"
    doc = Document()
    pdf = fitz.open(pdf_file)
    
    for page_num in range(len(pdf)):
        page = pdf[page_num]
        
        # Extract text
        text = page.get_text("text")
        doc.add_paragraph(text)
        
        # Extract images
        image_list = page.get_images()
        for img_index, img in enumerate(image_list):
            xref = img[0]
            base_image = pdf.extract_image(xref)
            image_bytes = base_image["image"]
            
            # Convert image bytes to PIL Image
            image = Image.open(io.BytesIO(image_bytes))
            
            # Save image temporarily
            temp_img_path = f"temp_img_{page_num}_{img_index}.png"
            image.save(temp_img_path)
            
            # Add image to document
            doc.add_picture(temp_img_path, width=Inches(6.0))
            
            # Clean up temporary image file
            import os
            os.remove(temp_img_path)
        
        # Add page break between pages
        if page_num < len(pdf) - 1:
            doc.add_page_break()
    
    doc.save(output_file)
    pdf.close()
    
    return output_file


def convert_to_image(pdf_file: str) -> List[str]:
    doc = fitz.open(pdf_file)
    output_files = []
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        zoom = 2.0
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        output_file = f"page_{page_num + 1}.png"
        pix.save(output_file)
        output_files.append(output_file)
    
    doc.close()
    return output_files

def convert_to_excel(pdf_file: str) -> str:
    output_file = "output.xlsx"
    doc = fitz.open(pdf_file)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Extracted Data"
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text()
        for line in text.split('\n'):
            if line.strip():
                ws.append([line])
    
    wb.save(output_file)
    doc.close()
    return output_file

def convert_to_epub(pdf_file: str) -> str:
    output_file = "output.epub"
    doc = fitz.open(pdf_file)
    book = epub.EpubBook()
    
    book.set_title("Converted PDF")
    book.set_language("en")
    
    content = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text()
        content.append(f"<div>{text}</div>")
    
    chapter = epub.EpubHtml(title="PDF Content", file_name="chapter1.xhtml")
    chapter.content = f"<h1>PDF Content</h1>{''.join(content)}"
    
    book.add_item(chapter)
    book.toc = [(epub.Section('PDF Content'), [chapter])]
    book.spine = ['nav', chapter]
    
    epub.write_epub(output_file, book)
    doc.close()
    return output_file
